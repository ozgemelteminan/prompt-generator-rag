"""Offline parser adapters for supported, validated document media types."""

import re
from io import BytesIO

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.document_processing.models import DocumentMetadata, ParsedDocument, TextBlock
from app.document_processing.parsing import DocumentParser, DocumentParserError

_MARKDOWN_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_MARKDOWN_LIST = re.compile(r"^\s*(?:[-*+] |\d+[.)] )(.+?)\s*$")
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class PdfDocumentParser:
    def parse(self, content: bytes, metadata: DocumentMetadata) -> ParsedDocument:
        try:
            reader = PdfReader(BytesIO(content))
            blocks: list[TextBlock] = []
            for page_number, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                for paragraph in _paragraphs(text):
                    blocks.append(
                        TextBlock("paragraph", paragraph, len(blocks), page_number=page_number)
                    )
            return ParsedDocument(metadata.document_id, tuple(blocks))
        except Exception as error:
            raise DocumentParserError("PDF parsing failed.") from error


class DocxDocumentParser:
    def parse(self, content: bytes, metadata: DocumentMetadata) -> ParsedDocument:
        try:
            document = DocxDocument(BytesIO(content))
            blocks: list[TextBlock] = []
            section: str | None = None
            for paragraph in document.paragraphs:
                text = paragraph.text
                if not text.strip():
                    continue
                style_name = paragraph.style.name if paragraph.style is not None else ""
                match = re.match(r"Heading\s+([1-9][0-9]*)$", style_name, flags=re.IGNORECASE)
                if match:
                    level = int(match.group(1))
                    section = text
                    blocks.append(TextBlock("heading", text, len(blocks), heading_level=level))
                else:
                    blocks.append(TextBlock("paragraph", text, len(blocks), section=section))
            return ParsedDocument(metadata.document_id, tuple(blocks))
        except Exception as error:
            raise DocumentParserError("DOCX parsing failed.") from error


class TextDocumentParser:
    def parse(self, content: bytes, metadata: DocumentMetadata) -> ParsedDocument:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError as error:
            raise DocumentParserError("Text decoding failed.") from error
        blocks = tuple(
            TextBlock("paragraph", paragraph, index)
            for index, paragraph in enumerate(_paragraphs(text))
        )
        return ParsedDocument(metadata.document_id, blocks)


class MarkdownDocumentParser:
    def parse(self, content: bytes, metadata: DocumentMetadata) -> ParsedDocument:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError as error:
            raise DocumentParserError("Markdown decoding failed.") from error
        blocks: list[TextBlock] = []
        paragraph: list[str] = []
        code: list[str] = []
        section: str | None = None
        in_code = False

        def flush_paragraph() -> None:
            if paragraph:
                blocks.append(
                    TextBlock("paragraph", "\n".join(paragraph), len(blocks), section=section)
                )
                paragraph.clear()

        for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
            if line.strip().startswith("```"):
                flush_paragraph()
                if in_code:
                    blocks.append(TextBlock("code", "\n".join(code), len(blocks), section=section))
                    code.clear()
                in_code = not in_code
                continue
            if in_code:
                code.append(line)
                continue
            heading = _MARKDOWN_HEADING.match(line)
            if heading:
                flush_paragraph()
                section = heading.group(2)
                blocks.append(
                    TextBlock("heading", section, len(blocks), heading_level=len(heading.group(1)))
                )
                continue
            list_item = _MARKDOWN_LIST.match(line)
            if list_item:
                flush_paragraph()
                blocks.append(
                    TextBlock("list_item", list_item.group(1), len(blocks), section=section)
                )
                continue
            if not line.strip():
                flush_paragraph()
            else:
                paragraph.append(line)
        flush_paragraph()
        if code:
            blocks.append(TextBlock("code", "\n".join(code), len(blocks), section=section))
        return ParsedDocument(metadata.document_id, tuple(blocks))


class DocumentParserRegistry:
    def __init__(self) -> None:
        self._parsers: dict[str, DocumentParser] = {
            "application/pdf": PdfDocumentParser(),
            DOCX_MEDIA_TYPE: DocxDocumentParser(),
            "text/plain": TextDocumentParser(),
            "text/markdown": MarkdownDocumentParser(),
        }

    def for_media_type(self, media_type: str) -> DocumentParser:
        try:
            return self._parsers[media_type]
        except KeyError as error:
            raise DocumentParserError("No parser is available for this document type.") from error


def _paragraphs(text: str) -> list[str]:
    return [paragraph for paragraph in re.split(r"\n\s*\n+", text) if paragraph.strip()]
